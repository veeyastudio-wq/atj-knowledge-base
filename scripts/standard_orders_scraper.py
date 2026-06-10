#!/usr/bin/env python3
"""
ATJ Knowledge Base — Category 5 Scraper
Standard Family Orders: Volume 1 (Financial Remedies) and Volume 2 (Children and Other).
Source: judiciary.uk — May 2026 release.

Downloads two zip archives and one standalone docx, extracts and parses each
.docx file using python-docx, and saves one markdown file per order to
raw/standard_orders/.

RAG design: one file per order document. Each order is already a discrete,
named document — splitting within them would break coherence.

Run from repo root:
    python scripts/standard_orders_scraper.py

Output: raw/standard_orders/
"""

import io
import logging
import os
import re
import sys
import time
import zipfile
from datetime import date
from typing import List, Optional, Tuple

import requests

try:
    from docx import Document
    HAS_DOCX = True
except ImportError:
    HAS_DOCX = False

# ---------------------------------------------------------------------------
# Config
# ---------------------------------------------------------------------------

OUTPUT_DIR = "raw/standard_orders"
LOG_FILE = os.path.join(OUTPUT_DIR, "_scrape_log.txt")
DELAY_SECONDS = 2
TODAY = date.today().isoformat()

SOURCE_PAGE = (
    "https://www.judiciary.uk/guidance-and-resources/"
    "update-from-mr-justice-peel-judge-in-charge-of-the-standard-orders/"
)

HEADERS = {
    "User-Agent": (
        "Mozilla/5.0 (compatible; ATJ-KB-Scraper/1.0; "
        "+https://github.com/veeyastudio-wq/atj-knowledge-base)"
    )
}

SOURCES = [
    {
        "slug": "sfo_vol1",
        "label": "Standard Family Orders — Volume 1: Financial Remedies (May 2026)",
        "url": "https://www.judiciary.uk/wp-content/uploads/2026/05/SFO-Volume-1-05.26.zip",
        "volume": "1",
        "is_zip": True,
    },
    {
        "slug": "sfo_vol2",
        "label": "Standard Family Orders — Volume 2: Children and Other Orders (May 2026)",
        "url": "https://www.judiciary.uk/wp-content/uploads/2026/05/SFO-Volume-2-05.26-1.zip",
        "volume": "2",
        "is_zip": True,
    },
    {
        "slug": "sfo_house_rules",
        "label": "Standard Family Orders — House Rules (May 2023)",
        "url": "https://www.judiciary.uk/wp-content/uploads/2023/05/HouseRules.May2023.docx",
        "volume": "house_rules",
        "is_zip": False,
    },
]

# ---------------------------------------------------------------------------
# Logging
# ---------------------------------------------------------------------------

def _setup_logging() -> logging.Logger:
    os.makedirs(OUTPUT_DIR, exist_ok=True)
    logger = logging.getLogger("standard_orders_scraper")
    logger.setLevel(logging.DEBUG)
    fmt = logging.Formatter("%(asctime)s [%(levelname)s] %(message)s")
    fh = logging.FileHandler(LOG_FILE, encoding="utf-8")
    fh.setFormatter(fmt)
    sh = logging.StreamHandler(sys.stdout)
    sh.setFormatter(fmt)
    logger.addHandler(fh)
    logger.addHandler(sh)
    return logger


log = _setup_logging()

# ---------------------------------------------------------------------------
# YAML helpers
# ---------------------------------------------------------------------------

_YAML_SPECIAL = re.compile(r'[:#\[\]{},&*?|<>=!%@`]')


def _yaml_str(value: str) -> str:
    if _YAML_SPECIAL.search(value):
        escaped = value.replace('"', '\\"')
        return f'"{escaped}"'
    return value


def _frontmatter(fields: dict) -> str:
    lines = ["---"]
    for k, v in fields.items():
        if v is None:
            lines.append(f"{k}:")
        elif isinstance(v, list):
            lines.append(f"{k}:")
            for item in v:
                lines.append(f"  - {_yaml_str(str(item))}")
        else:
            lines.append(f"{k}: {_yaml_str(str(v))}")
    lines.append("---\n")
    return "\n".join(lines)

# ---------------------------------------------------------------------------
# Slug helpers
# ---------------------------------------------------------------------------

def _slugify(text: str) -> str:
    text = text.lower().strip()
    text = re.sub(r"[^\w\s-]", "", text)
    text = re.sub(r"[\s_-]+", "_", text)
    text = re.sub(r"^_+|_+$", "", text)
    return text[:80]


def _order_number_from_filename(filename: str) -> str:
    """
    Extract order number from docx filename.
    e.g. 'Order-1.1.docx' -> '1.1'
         'Order-7.0-Financial-Remedy.docx' -> '7.0'
         'HouseRules.May2023.docx' -> 'house_rules'
    """
    stem = os.path.splitext(os.path.basename(filename))[0]
    match = re.search(r'[Oo]rder[-_](\d+\.\d+)', stem)
    if match:
        return match.group(1)
    return _slugify(stem)

# ---------------------------------------------------------------------------
# HTTP session
# ---------------------------------------------------------------------------

def _make_session() -> requests.Session:
    session = requests.Session()
    session.headers.update(HEADERS)
    return session

# ---------------------------------------------------------------------------
# docx parsing
# ---------------------------------------------------------------------------

def _parse_docx_bytes(data: bytes, filename: str) -> str:
    """
    Parse a docx file from bytes and return plain text.
    Preserves paragraph structure. Skips empty paragraphs.
    """
    try:
        doc = Document(io.BytesIO(data))
    except Exception as e:
        log.warning(f"  Could not parse {filename}: {e}")
        return ""

    paragraphs = []
    for para in doc.paragraphs:
        text = para.text.strip()
        if text:
            paragraphs.append(text)

    # Also extract text from tables
    for table in doc.tables:
        for row in table.rows:
            row_cells = []
            for cell in row.cells:
                cell_text = cell.text.strip()
                if cell_text:
                    row_cells.append(cell_text)
            if row_cells:
                paragraphs.append(" | ".join(row_cells))

    return "\n\n".join(paragraphs)


def _docx_title(data: bytes, filename: str) -> str:
    """
    Attempt to extract a meaningful title from the docx.
    Falls back to the filename stem.
    """
    try:
        doc = Document(io.BytesIO(data))
        if doc.core_properties.title:
            return doc.core_properties.title.strip()
        for para in doc.paragraphs:
            text = para.text.strip()
            if text and len(text) > 3:
                return text[:120]
    except Exception:
        pass
    return os.path.splitext(os.path.basename(filename))[0]

# ---------------------------------------------------------------------------
# Fetch and save
# ---------------------------------------------------------------------------

def _fetch_bytes(session: requests.Session, url: str) -> Optional[bytes]:
    log.info(f"Fetching: {url}")
    try:
        resp = session.get(url, timeout=60, stream=True)
        resp.raise_for_status()
        time.sleep(DELAY_SECONDS)
        return resp.content
    except requests.RequestException as e:
        log.error(f"Failed to fetch {url}: {e}")
        return None


def _save_order(
    docx_data: bytes,
    filename: str,
    idx: int,
    source: dict,
) -> bool:
    order_number = _order_number_from_filename(filename)
    title = _docx_title(docx_data, filename)
    body = _parse_docx_bytes(docx_data, filename)

    if not body:
        log.warning(f"  No text extracted from {filename} — skipping")
        return False

    volume = source["volume"]
    file_slug = _slugify(order_number)
    out_filename = f"sfo_vol{volume}__{idx:03d}__{file_slug}.md"

    if volume == "house_rules":
        out_filename = f"sfo_house_rules__000__house_rules.md"

    filepath = os.path.join(OUTPUT_DIR, out_filename)

    fm = _frontmatter({
        "source": source["label"],
        "title": title,
        "order_number": order_number,
        "volume": volume,
        "parent_document": source["label"],
        "parent_slug": source["slug"],
        "section_index": idx,
        "source_filename": os.path.basename(filename),
        "source_page": SOURCE_PAGE,
        "url": source["url"],
        "scrape_date": TODAY,
        "licence": "Open Government Licence v3.0",
        "content_type": "standard_order",
        "jurisdiction": "England and Wales",
    })

    with open(filepath, "w", encoding="utf-8") as f:
        f.write(fm)
        f.write(f"# {title}\n\n")
        f.write(body)
        f.write("\n")

    log.info(f"  Saved: {out_filename}")
    return True


def process_zip_source(session: requests.Session, source: dict) -> int:
    data = _fetch_bytes(session, source["url"])
    if not data:
        return 0

    try:
        zf = zipfile.ZipFile(io.BytesIO(data))
    except zipfile.BadZipFile as e:
        log.error(f"Bad zip for {source['slug']}: {e}")
        return 0

    docx_names = sorted([
        n for n in zf.namelist()
        if n.lower().endswith(".docx")
        and not os.path.basename(n).startswith("~")
        and not os.path.basename(n).startswith(".")
    ])

    log.info(f"  Found {len(docx_names)} docx files in zip")

    saved = 0
    for idx, name in enumerate(docx_names):
        time.sleep(DELAY_SECONDS)
        try:
            docx_data = zf.read(name)
        except Exception as e:
            log.warning(f"  Could not read {name} from zip: {e}")
            continue

        ok = _save_order(docx_data, name, idx, source)
        if ok:
            saved += 1

    return saved


def process_docx_source(session: requests.Session, source: dict) -> int:
    data = _fetch_bytes(session, source["url"])
    if not data:
        return 0

    ok = _save_order(data, source["url"], 0, source)
    return 1 if ok else 0

# ---------------------------------------------------------------------------
# Main
# ---------------------------------------------------------------------------

def main() -> None:
    if not HAS_DOCX:
        log.error(
            "python-docx is not installed. "
            "Run: pip install python-docx --break-system-packages"
        )
        sys.exit(1)

    os.makedirs(OUTPUT_DIR, exist_ok=True)
    session = _make_session()
    total = 0

    log.info("=== Category 5 scraper starting ===")
    log.info(f"Output: {OUTPUT_DIR}/")

    for source in SOURCES:
        log.info(f"\n--- {source['label']} ---")
        if source["is_zip"]:
            count = process_zip_source(session, source)
        else:
            count = process_docx_source(session, source)
        log.info(f"{source['slug']}: {count} files saved")
        total += count

    log.info(f"\n=== Done. {total} files saved to {OUTPUT_DIR}/ ===")


if __name__ == "__main__":
    main()
